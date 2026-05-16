"""Replace the Background of the Study section in FINAL THESIS 3 PAPER.docx.

Strategy:
- Heading at paragraph index 139 stays untouched.
- Body paragraphs 140-157 get overwritten with the new 8-paragraph content
  (interleaved with blank paragraphs).
- The next heading "Significance of the Study" at index 158 stays untouched.
"""
from copy import deepcopy
from docx import Document

NEW_PARAGRAPHS = [
    "Modern software development has reached a scale where the cost of code that is hard to read has become measurable in trillions of dollars. The Consortium for IT Software Quality estimated that poor software quality cost United States organizations approximately USD 2.41 trillion in 2022, with the bulk of that figure traced to unaddressed technical debt and defects that survived review (Krasner, 2022). The Cost of a Data Breach Report further found that breaches attributable to “system complexity” — code that cannot be audited quickly — added an average of USD 241,000 per incident (IBM Security, 2023). These figures reframe code readability from a stylistic preference into a measurable economic concern, and they place the cost of unreadable software in the same conversation as missed deadlines and security failures rather than separate from them.",

    "Software development today is also more collaborative and incremental than it has ever been. Contributors are routinely asked not only to write code but to understand, maintain, explain, and extend systems written by others (Rukmono et al., 2021; Romeo et al., 2024). The growing adoption of artificial intelligence assistants has intensified this trend. GitHub (2024) reported that developers using AI coding assistants ship code approximately 55 percent faster, but the same studies show that reviewers spend correspondingly more time per pull request, because the comprehension burden simply shifts from authoring to reading. Stack Overflow’s 2024 Developer Survey similarly found that while 76 percent of developers use or plan to use AI tools, 45 percent report that they do not fully trust the output. The bottleneck has moved from writing to comprehension, and the ability to interpret unfamiliar code quickly has become a defining productivity skill rather than a peripheral one.",

    "One of the most established responses to the comprehension problem is the use of software design patterns. Originally codified by Gamma, Helm, Johnson, and Vlissides (1994), design patterns provide reusable solutions, common structures, and a shared vocabulary for object-oriented software design. A systematic literature review spanning more than two decades of empirical studies confirmed that classes implementing recognized design patterns exhibit measurably lower change-proneness than ad-hoc structures of comparable size (Ampatzoglou, Charalampidou, & Stamelos, 2020), and Nesteruk (2022) demonstrated that the original Gang-of-Four arrangements translate cleanly into modern C++20 idioms. Knowledge of design patterns supports code review, system maintenance, the communication of design intent, and the interpretation of existing software structures — making pattern literacy one of the more durable forms of engineering knowledge a contributor can acquire.",

    "For interns and novice developers preparing for professional software development environments, design-pattern literacy provides a vocabulary that compresses long structural explanations into single names. A learner who recognizes a Singleton, a Strategy, or an Adapter in a real codebase can grasp the intent of an entire component at a glance, allowing them to participate in code reviews, ask precise questions, and contribute meaningfully far earlier in the onboarding cycle than a learner who must reconstruct intent line by line. For experienced programmers, the same vocabulary functions as a compression layer over architecture discussions: pattern names let teams negotiate design decisions in minutes rather than hours, reduce miscommunication during handovers, and provide a stable reference that survives team turnover. In both cases the benefit is the same — pattern literacy converts implicit structural knowledge into transferable, reviewable speech.",

    "Despite these benefits, design patterns remain difficult for learners to recognize in actual source code. Interns and novice developers may correctly define patterns in the abstract, yet still struggle to identify the same patterns when reading real C++ implementations. Pattern evidence in production code is distributed across class declarations, methods, attributes, object-creation logic, inheritance, composition, and the relationships among components, and recognizing it requires a kind of structural reading that conceptual lessons alone do not produce. Documentation has historically been expected to bridge this gap, but manual documentation is often incomplete, delayed, or inconsistent with the implementation it describes, and the resulting drift leaves learners without a reliable reference (Romeo et al., 2024).",

    "The theory-to-practice gap is most damaging in organizations that exhibit a particular profile. These are organizations with absent or under-documented coding standards, a reliance on informal mentorship rather than written design conventions, rapid contributor turnover, mixed-skill teams that co-locate novices and experts without a shared vocabulary, and limited code-review bandwidth among senior developers (Steinmacher et al., 2021; Storey et al., 2021; Aniche, Treude, & Bacchelli, 2022). Organizations matching this profile include open-source communities, volunteer-driven technology non-profits, bootcamp-led teams, internship programs, and small-to-medium development teams with high turnover. In each of these settings, the burden of converting design-pattern theory into practical recognition falls almost entirely on the new contributor and on whoever happens to be available to mentor them. This produces a bottleneck that scales poorly and that does not persist reliably from one onboarding cycle to the next, because knowledge transferred informally during one cohort tends to leave the organization with that cohort.",

    "DEVCON Luzon, a Filipino non-profit technology community operating through chapter-based internships and bootcamps, exhibits each of the qualities described above and serves as the primary evaluation context of this study. Its volunteer-driven structure, intermittent mentorship availability, and recurring intake of interns from different academic backgrounds make it representative of the broader organizational profile rather than an exception to it. Findings from the DEVCON Luzon evaluation context are therefore intended to generalize to organizations that share this profile, not solely to DEVCON Luzon itself, and the qualitative and quantitative results reported in later chapters should be read with that scope in mind.",

    "To address the comprehension and onboarding challenges described above, this study developed CodiNeo, a documentation generation and design-pattern learning support system for C++ programming. CodiNeo provides learning modules that introduce selected software design-pattern concepts and a C++ analysis studio that examines source code, detects supported design-pattern evidence, and generates documentation-oriented outputs that help users connect theoretical knowledge with actual implementation. The system is grounded in structure-aware source-code analysis: instead of treating code as plain text alone, CodiNeo examines class-level structures, methods, attributes, and relationships that may serve as evidence for supported design patterns. The detailed catalog of supported patterns and the literature behind each is presented in Chapter 2, while the algorithmic mechanics of structural analysis, evidence detection, and documentation generation are presented in Chapter 3. The remainder of this chapter establishes the study’s significance, statement of the problem, scope, and definitions.",
]


def _copy_run_formatting(template_run, new_run):
    """Best-effort copy of font formatting from one run to another."""
    try:
        new_run.bold = template_run.bold
        new_run.italic = template_run.italic
        new_run.underline = template_run.underline
        if template_run.font and template_run.font.name:
            new_run.font.name = template_run.font.name
        if template_run.font and template_run.font.size:
            new_run.font.size = template_run.font.size
    except Exception:
        pass


def overwrite_paragraph(paragraph, text, template_run=None):
    """Clear all runs in a paragraph and write a single new run with text."""
    # Remove existing runs at the XML level.
    p_elem = paragraph._element
    for r in list(p_elem.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')):
        p_elem.remove(r)
    new_run = paragraph.add_run(text)
    if template_run is not None:
        _copy_run_formatting(template_run, new_run)


def main():
    src = 'FINAL THESIS 3 PAPER.docx'
    d = Document(src)

    # Sanity-check anchors.
    assert d.paragraphs[139].text.strip() == 'Background of the Study', \
        f'Unexpected heading at 139: {d.paragraphs[139].text!r}'
    assert d.paragraphs[158].text.strip() == 'Significance of the Study', \
        f'Unexpected heading at 158: {d.paragraphs[158].text!r}'

    # Pick a template run from the existing body for formatting reuse.
    template_run = None
    for i in range(140, 158):
        runs = d.paragraphs[i].runs
        if runs and runs[0].text.strip():
            template_run = runs[0]
            break

    # Body slots: paragraphs 140..157 inclusive = 18 slots.
    # Plan: prose, blank, prose, blank, ... => 8 prose + 7 blanks = 15 used,
    # 3 trailing blanks remain.
    slots = list(range(140, 158))
    plan = []
    for idx, para_text in enumerate(NEW_PARAGRAPHS):
        plan.append(para_text)
        if idx != len(NEW_PARAGRAPHS) - 1:
            plan.append('')  # blank spacer
    # Pad to fill remaining slots with blank paragraphs.
    while len(plan) < len(slots):
        plan.append('')

    assert len(plan) == len(slots), f'plan/slots mismatch {len(plan)} vs {len(slots)}'

    for slot_idx, text in zip(slots, plan):
        overwrite_paragraph(d.paragraphs[slot_idx], text, template_run)

    out = 'FINAL THESIS 3 PAPER.docx'
    d.save(out)
    print('Saved', out)


if __name__ == '__main__':
    main()
