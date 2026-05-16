"""
Apply v3 SOP/SO rewrites to FINAL-THESIS-3-PAPER-1 (2).docx.

The target paragraphs (verified by scan):
  SOPs : indices 174..178  (List Paragraph style, under "Statement of the Problem")
  SOs  : indices 182..186  (List Paragraph style, under "Objectives of the Study")

v3 rewrites are sourced verbatim from _thesis_work3/isolation-audit.md blocks A and B.
v3 removes "C++" from all SOPs and from SOs 1, 3, 4, 5 (SO 2 keeps C++ only as the
prototype's first language with a cross-reference to Section 1.7 Scope), and removes
"DEVCON Luzon" entirely from SOPs and SOs.

Sanity check strategy:
  The previous version of this script checked that each target paragraph *starts*
  with the ORIGINAL vague wording. That worked for a first run but blocks re-runs
  after v2 was applied. v3's sanity check uses CONTEXT ANCHORS instead:
    - paragraph index falls inside the expected range
    - the SOP target range sits between the "Statement of the Problem" heading
      (¶171) and the "Objectives of the Study" heading (¶180)
    - the SO target range sits between "Objectives of the Study" (¶180) and
      "Theoretical Framework" (¶188)
    - every target paragraph has style "List Paragraph"
  This makes the script idempotent: running it twice in a row writes the same
  v3 text both times without errors.

Safety:
  Backup file already exists:
    FINAL-THESIS-3-PAPER-1 (2).backup-before-sop-rewrite.docx
  That backup pre-dates any rewrite (v2 or v3) and remains the rollback path.
"""

from pathlib import Path
from typing import Final, Mapping

import docx
from docx.document import Document

DOC_PATH: Final[Path] = Path("FINAL-THESIS-3-PAPER-1 (2).docx")

# Context-anchor headings used by the sanity check.
SOP_HEADING_TEXT: Final[str] = "Statement of the Problem"
SO_HEADING_TEXT: Final[str] = "Objectives of the Study"
NEXT_HEADING_TEXT: Final[str] = "Theoretical Framework"

# Expected paragraph index ranges (inclusive on both ends).
SOP_RANGE: Final[range] = range(174, 179)
SO_RANGE: Final[range] = range(182, 187)

# Index -> new text. v3 rewrites verbatim from isolation-audit.md blocks A and B.
NEW_TEXT: Final[Mapping[int, str]] = {
    # --- Statement of the Problem (v3) ---
    174: (
        "How can the system teach design patterns in a way that lets the user "
        "try each pattern on their own source code immediately after the "
        "lesson, instead of only reading about it?"
    ),
    175: (
        "How can the system examine a user's source code to find which design "
        "patterns are present, without changing the user's original code, and "
        "allow the team to add support for new patterns later without "
        "rebuilding the system?"
    ),
    176: (
        "How can the system present its findings in a way that shows the user "
        "which lines of their code triggered each detected pattern and why, "
        "so the user can verify the finding for themselves?"
    ),
    177: (
        "How can the system write per-class documentation that combines "
        "findings the system verified from the user's code with explanations "
        "written by AI, while clearly labelling which sentences come from "
        "which source so the user can trust the documentation?"
    ),
    178: (
        "How well does the system help novice developers identify, document, "
        "and explain design patterns in source code that someone else wrote "
        "— measured by comparing what they can do before and after using the "
        "system?"
    ),
    # --- Specific Objectives (v3) ---
    182: (
        "Create learning modules where each design-pattern lesson is paired "
        "with a hands-on check the system runs on the user's own source "
        "code, so the user sees the lesson applied to their work right after "
        "learning it."
    ),
    183: (
        "Build a code-analysis feature that reads the user's source without "
        "modifying it, identifies which design patterns appear at the class "
        "level, and reads its pattern definitions from configuration files "
        "so new patterns can be added by editing files instead of rebuilding "
        "the system. The prototype implementation supports C++ as its first "
        "and only currently-supported language; extension to additional "
        "statically-typed object-oriented languages is identified as future "
        "work (see Section 1.7 Scope)."
    ),
    184: (
        "Build a results view that lists the most likely pattern matches "
        "first, links each match to the exact lines of code that caused it, "
        "and explains in plain language what the system saw in those lines."
    ),
    185: (
        "Build a documentation feature that produces a written explanation "
        "for each analysed class, mixing facts the system verified against "
        "the user's code with sentences written by AI, with each AI-written "
        "sentence visibly marked so the user can tell the two apart."
    ),
    186: (
        "Evaluate the system with novice-developer participants by "
        "measuring, before and after they use the system, how accurately "
        "they can identify design patterns in unfamiliar source code, how "
        "completely they can document those patterns, and how clearly they "
        "can explain why each pattern is there."
    ),
}


def replace_paragraph_text(paragraph: object, new_text: str) -> None:
    """Replace all text in `paragraph` with `new_text`, preserving the
    paragraph's style (List Paragraph numbering) and the first run's font
    properties.
    """
    runs = paragraph.runs  # type: ignore[attr-defined]
    template_run = runs[0] if runs else None

    for run in list(runs):
        run._element.getparent().remove(run._element)

    new_run = paragraph.add_run(new_text)  # type: ignore[attr-defined]

    if template_run is not None:
        new_run.font.name = template_run.font.name
        new_run.font.size = template_run.font.size
        new_run.font.bold = template_run.font.bold
        new_run.font.italic = template_run.font.italic
        new_run.font.underline = template_run.font.underline
        if template_run.font.color and template_run.font.color.rgb:
            new_run.font.color.rgb = template_run.font.color.rgb


def find_heading_index(doc: Document, heading_text: str) -> int:
    """Return the paragraph index of the first paragraph whose stripped text
    equals `heading_text`. Raises SystemExit if not found.
    """
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text:
            return i
    raise SystemExit(f"ABORT: heading '{heading_text}' not found in docx")


def check_context_anchors(doc: Document) -> None:
    """Validate that the SOP/SO target ranges still sit between the expected
    headings and that every target paragraph has style 'List Paragraph'.
    """
    sop_heading_idx = find_heading_index(doc, SOP_HEADING_TEXT)
    so_heading_idx = find_heading_index(doc, SO_HEADING_TEXT)
    next_heading_idx = find_heading_index(doc, NEXT_HEADING_TEXT)

    problems: list[str] = []

    if not (sop_heading_idx < SOP_RANGE.start <= SOP_RANGE.stop - 1 < so_heading_idx):
        problems.append(
            f"  SOP range {list(SOP_RANGE)} does not fall between "
            f"'{SOP_HEADING_TEXT}' (¶{sop_heading_idx}) and "
            f"'{SO_HEADING_TEXT}' (¶{so_heading_idx})"
        )

    if not (so_heading_idx < SO_RANGE.start <= SO_RANGE.stop - 1 < next_heading_idx):
        problems.append(
            f"  SO range {list(SO_RANGE)} does not fall between "
            f"'{SO_HEADING_TEXT}' (¶{so_heading_idx}) and "
            f"'{NEXT_HEADING_TEXT}' (¶{next_heading_idx})"
        )

    for idx in list(SOP_RANGE) + list(SO_RANGE):
        if idx >= len(doc.paragraphs):
            problems.append(f"  paragraph ¶{idx} is out of range")
            continue
        style_name = doc.paragraphs[idx].style.name
        if style_name != "List Paragraph":
            problems.append(
                f"  ¶{idx} has style '{style_name}', expected 'List Paragraph'"
            )

    if problems:
        print("ABORT: context anchors do not match expectations. No edits made.")
        for p in problems:
            print(p)
        raise SystemExit(1)


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"docx not found: {DOC_PATH.resolve()}")

    doc = docx.Document(str(DOC_PATH))

    check_context_anchors(doc)

    for idx, new_text in NEW_TEXT.items():
        replace_paragraph_text(doc.paragraphs[idx], new_text)
        print(f"  ¶{idx} rewrote -> {new_text[:80]}...")

    doc.save(str(DOC_PATH))
    print(f"\nSaved: {DOC_PATH}")
    print("Backup: FINAL-THESIS-3-PAPER-1 (2).backup-before-sop-rewrite.docx")


if __name__ == "__main__":
    main()
