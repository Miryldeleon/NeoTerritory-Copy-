"""
Apply v3 thesis edits beyond the SOP/SO rewrites:

  1. Scope addition (block C from _thesis_work3/isolation-audit.md): insert two
     new paragraphs into Section 1.7 Scope, immediately after the first body
     paragraph and before the existing "The system includes a learner pathway"
     paragraph. The two new paragraphs own BOTH the language boundary
     (statically-typed OO, C++ as prototype, dynamic-OO out of scope) and the
     evaluation-site boundary (DEVCON Luzon as primary evaluation context).

  2. Definition of Terms fill-in (block D): replace the placeholder entry for
     "Hash-Based Virtual Structural Copy" in Section 1.8. The placeholder
     currently reads `Hash-Based Virtual Structural Copy – This is your "star"
     term...`. The replacement matches the surrounding "Term – Definition"
     pattern with em-dash separator.

  3. (Optional) Heading-style promotion for "Background of the Study" (¶139)
     and "Significance of the Study" (¶156): change paragraph style from
     "Normal" to "Heading 2" so they appear in Word's Navigation Pane and any
     regenerated Table of Contents. Brings them in line with the other
     Chapter 1 section headings ("Statement of the Problem" ¶171, "Objectives
     of the Study" ¶180, "Scope and Delimitations" ¶227, "Definition of Terms"
     ¶261) which are already Heading 2.

Sanity check strategy:
  - Locate target paragraphs by stable textual anchors (heading names + the
    placeholder's opening words), not by hard-coded indices. The previous
    SOP/SO script anchors by index; Scope and DoT use text anchors so this
    script is robust to upstream paragraph insertions.
  - Scope insertion is idempotent: detects if the v3 paragraphs are already
    present (by their opening phrase "The algorithm at the core of this
    study") and skips re-insertion.
  - DoT replacement is idempotent: writes the v3 text whether the line
    currently holds the placeholder or already holds the v3 text.

Backup:
  FINAL-THESIS-3-PAPER-1 (2).backup-before-sop-rewrite.docx pre-dates every
  rewrite and remains the rollback path.
"""

from copy import deepcopy
from pathlib import Path
from typing import Final

import docx
from docx.document import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOC_PATH: Final[Path] = Path("FINAL-THESIS-3-PAPER-1 (2).docx")

# -- Scope ------------------------------------------------------------------

SCOPE_HEADING_TEXT: Final[str] = "Scope"
SCOPE_FIRST_PARA_PREFIX: Final[str] = "This study focuses on the design"
SCOPE_NEXT_PARA_PREFIX: Final[str] = "The system includes a learner pathway"

SCOPE_LANGUAGE_PARA: Final[str] = (
    "The algorithm at the core of this study is designed for statically-typed "
    "object-oriented programming languages with class-level syntactic "
    "structure, including but not limited to C++, Java, C#, and Kotlin. The "
    "prototype implementation supports C++ as its first and only "
    "currently-supported language, chosen because of the time and resource "
    "constraints of an undergraduate thesis; extension to additional "
    "statically-typed object-oriented languages is identified as future work "
    "in Chapter 5. Support for dynamically-typed object-oriented languages "
    "(e.g. Python, JavaScript) is out of scope for both the prototype and "
    "the algorithmic claim of this study, and is not committed to as future "
    "work."
)

SCOPE_EVALUATION_PARA: Final[str] = (
    "The primary evaluation context of this study is the DEVCON Luzon "
    "learning environment, where the system is tested with novice-developer "
    "participants drawn from the DEVCON Luzon community. Findings reported "
    "in Chapters 4 and 5 describe usefulness and effectiveness within that "
    "participant pool and are not intended to generalize beyond it. The "
    "system itself is built without DEVCON-Luzon-specific assumptions — the "
    "partnership defines the evaluation site, not the audience the system "
    "is technically applicable to. Industry-scale rationale for the broader "
    "audiences for whom the system is technically applicable is discussed "
    "in the Significance section (Section 1.4)."
)

SCOPE_V3_MARKER: Final[str] = "The algorithm at the core of this study"

# -- Definition of Terms -----------------------------------------------------

DOT_PLACEHOLDER_PREFIX: Final[str] = "Hash-Based Virtual Structural Copy"

DOT_V3_TEXT: Final[str] = (
    "Hash-Based Virtual Structural Copy – An immutable representation of "
    "the user's source-code parse tree (the actual tree) is mirrored into "
    "a working copy (the virtual tree) on which classification tags, "
    "cross-references, and pattern-detection results are written. "
    "Structural hashing identifies repeated sub-tree shapes so that "
    "detection work is performed once per unique structural shape rather "
    "than per textual occurrence. The actual tree is never mutated; it "
    "provides an auditable ground truth to which every detected-pattern "
    "claim is anchored. This differs from a standard deep copy by "
    "maintaining structural identity through hashing rather than "
    "node-by-node duplication, and from a reference copy by ensuring "
    "writes to the working surface never propagate back to the source "
    "representation. The algorithm is defined for statically-typed "
    "object-oriented programming languages with class-level syntactic "
    "structure; the prototype implementation in this study operates on "
    "C++ parse trees (see Section 1.7 Scope for the full language-scope "
    "statement)."
)

# -- Heading promotion ------------------------------------------------------

PROMOTE_HEADINGS_TO_H2: Final[tuple[str, ...]] = (
    "Background of the Study",
    "Significance of the Study",
)


def find_paragraph_by_text(
    doc: Document, target_text: str, *, prefix_match: bool = False
) -> int:
    """Return the index of the first paragraph whose stripped text matches.

    When prefix_match=True, the paragraph's stripped text must START WITH
    target_text. When False (default), the match must be exact.
    Raises SystemExit if not found.
    """
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if prefix_match:
            if text.startswith(target_text):
                return i
        else:
            if text == target_text:
                return i
    needle = f"starts-with '{target_text}'" if prefix_match else f"'{target_text}'"
    raise SystemExit(f"ABORT: paragraph matching {needle} not found in docx")


def insert_paragraph_after(
    anchor: Paragraph, text: str
) -> Paragraph:
    """Insert a new paragraph immediately after `anchor`, with the same
    paragraph properties (style + numbering) as the anchor. Returns the new
    Paragraph wrapper.
    """
    # Clone the anchor's <w:p> element so paragraph properties (pPr) carry
    # over, then strip any runs from the clone.
    new_p_element = deepcopy(anchor._p)
    for run_el in new_p_element.findall(qn("w:r")):
        new_p_element.remove(run_el)
    anchor._p.addnext(new_p_element)

    new_paragraph = Paragraph(new_p_element, anchor._parent)
    new_paragraph.add_run(text)
    return new_paragraph


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Replace all runs in `paragraph` with a single run holding new_text.
    Preserves paragraph style + first-run font properties.
    """
    runs = paragraph.runs
    template_run = runs[0] if runs else None
    for run in list(runs):
        run._element.getparent().remove(run._element)
    new_run = paragraph.add_run(new_text)
    if template_run is not None:
        new_run.font.name = template_run.font.name
        new_run.font.size = template_run.font.size
        new_run.font.bold = template_run.font.bold
        new_run.font.italic = template_run.font.italic
        if template_run.font.color and template_run.font.color.rgb:
            new_run.font.color.rgb = template_run.font.color.rgb


def apply_scope_addition(doc: Document) -> None:
    """Insert the two v3 Scope paragraphs after the first scope body para.
    Idempotent: skips if the v3 marker phrase is already present.
    """
    # Skip if already applied.
    for p in doc.paragraphs:
        if p.text.strip().startswith(SCOPE_V3_MARKER):
            print("  scope: v3 paragraphs already present, skipping insertion")
            return

    first_para_idx = find_paragraph_by_text(
        doc, SCOPE_FIRST_PARA_PREFIX, prefix_match=True
    )
    anchor = doc.paragraphs[first_para_idx]
    print(f"  scope: anchor paragraph ¶{first_para_idx} = '{anchor.text[:60]}...'")

    # Insert in reverse order so they end up in forward order after the anchor.
    insert_paragraph_after(anchor, SCOPE_EVALUATION_PARA)
    insert_paragraph_after(anchor, SCOPE_LANGUAGE_PARA)

    print("  scope: inserted 2 new paragraphs (language + evaluation boundaries)")


def apply_dot_fillin(doc: Document) -> None:
    """Replace the Hash-Based Virtual Structural Copy placeholder with the v3
    definition. Idempotent: writes v3 text whether the line is the placeholder
    or already holds v3 content.
    """
    idx = find_paragraph_by_text(
        doc, DOT_PLACEHOLDER_PREFIX, prefix_match=True
    )
    target = doc.paragraphs[idx]
    print(f"  dot: target paragraph ¶{idx} = '{target.text[:80]}...'")
    replace_paragraph_text(target, DOT_V3_TEXT)
    print(f"  dot: replaced with v3 definition")


def apply_heading_promotion(doc: Document) -> None:
    """Promote the listed paragraph headings from Normal to Heading 2 so they
    appear in Word's Navigation Pane and any regenerated TOC.
    """
    for heading_text in PROMOTE_HEADINGS_TO_H2:
        try:
            idx = find_paragraph_by_text(doc, heading_text)
        except SystemExit as e:
            print(f"  heading: skip — {e}")
            continue
        para = doc.paragraphs[idx]
        old_style = para.style.name
        if old_style == "Heading 2":
            print(f"  heading: ¶{idx} '{heading_text}' already Heading 2, skipping")
            continue
        para.style = doc.styles["Heading 2"]
        print(
            f"  heading: ¶{idx} '{heading_text}' promoted "
            f"{old_style} -> Heading 2"
        )


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"docx not found: {DOC_PATH.resolve()}")

    doc = docx.Document(str(DOC_PATH))

    print("[1/3] Scope addition (Section 1.7)")
    apply_scope_addition(doc)

    print("\n[2/3] Definition of Terms fill-in (Section 1.8)")
    apply_dot_fillin(doc)

    print("\n[3/3] Heading promotion (Background + Significance)")
    apply_heading_promotion(doc)

    doc.save(str(DOC_PATH))
    print(f"\nSaved: {DOC_PATH}")


if __name__ == "__main__":
    main()
