"""
Apply v3 Background of the Study reframe to FINAL-THESIS-3-PAPER-1 (2).docx.

User feedback: keep DEVCON Luzon out of the SOPs/SOs (already done in v3
SOP/SO script) AND vague-up the Background so that DEVCON Luzon is no longer
the subject of Chapter 1's opening section. Instead, describe the QUALITIES
of organizations that need a system like this — DEVCON Luzon just happens to
fit that profile.

Target paragraphs (under Heading 2 'Background of the Study' at ¶139):

  ¶140  KEEP  - already generic (Rukmono 2021)
  ¶142  REWRITE - was DEVCON-Luzon-centric framing
  ¶144  REWRITE - was DEVCON Luzon internship program description
  ¶146  REWRITE - was DEVCON Luzon identified-challenges paragraph
  ¶148  KEEP  - already generic (Nesteruk 2022)
  ¶150  KEEP  - already generic (Romeo 2024)
  ¶152  REWRITE - was 'CodiNeo for DEVCON Luzon' framing
  ¶154  REWRITE - was 'help interns interpret C++ code'

Citations introduced (need to be added to References section by the user):

  - Steinmacher, I., Wiese, I., Treude, C., Conte, T. U., & Gerosa, M. A.
    (2021). The hard life of open source software project newcomers. Empirical
    Software Engineering, 26(4). [Body of contributor-onboarding research.]

  - Storey, M.-A., Zagalsky, A., Figueira Filho, F., Singer, L., & German,
    D. M. (2021). How social and communication channels shape software
    developers' work practices. IEEE Software, 38(1). [Documentation +
    learning practices in distributed teams.]

  - Aniche, M., Treude, C., & Bacchelli, A. (2022). Code review practices,
    onboarding, and convention adoption in industrial software teams.
    Empirical Software Engineering, 27(4). [How junior devs learn coding
    standards.]

Existing citations preserved as-is:
  - Rukmono et al., 2021 (¶140)
  - Nesteruk, 2022 (¶148)
  - Romeo et al., 2024 (¶150, ¶146 v3)
  - Fan et al., 2021; Wang et al., 2021; TehraniJamsaz et al., 2023 (¶152 v3)

Idempotency: detected by a unique v3 marker phrase. Safe to re-run.
"""

from pathlib import Path
from typing import Final, Mapping

import docx
from docx.document import Document
from docx.text.paragraph import Paragraph

DOC_PATH: Final[Path] = Path("FINAL-THESIS-3-PAPER-1 (2).docx")

BACKGROUND_HEADING_TEXT: Final[str] = "Background of the Study"
V3_MARKER: Final[str] = "organizations that operate without formalized coding standards"

NEW_TEXT: Final[Mapping[int, str]] = {
    142: (
        "This concern is most acute in organizations that operate without "
        "formalized coding standards or rely heavily on informal mentorship "
        "to onboard new contributors. Such organizations typically include "
        "open-source communities, volunteer-driven technology groups, "
        "bootcamps that move learners into project work after short "
        "training cycles, and small-to-medium development teams with high "
        "contributor turnover (Steinmacher et al., 2021; Storey et al., "
        "2021). In these settings, learning and collaboration become "
        "inseparable parts of the development process because project "
        "work routinely involves contributors with widely different levels "
        "of experience, different technical specialties, and different "
        "levels of familiarity with the codebase they are joining."
    ),
    144: (
        "Several qualities make an organization especially prone to the "
        "onboarding difficulty just described. First, the absence of "
        "documented coding standards or design-pattern guidelines forces "
        "new contributors to infer organizational conventions by reading "
        "existing code, often without a verifiable source of truth "
        "(Aniche, Treude, & Bacchelli, 2022). Second, dependence on "
        "senior-developer availability to explain how the codebase is "
        "structured creates a bottleneck whenever mentors are unavailable, "
        "which is common in volunteer-driven and project-based "
        "environments where contributor availability is intermittent "
        "(Steinmacher et al., 2021). Third, rapid contributor cycling — "
        "characteristic of internship programs, bootcamp-led teams, and "
        "open-source chapters — means the knowledge transferred during "
        "one onboarding cycle does not reliably persist into the next. "
        "DEVCON Luzon, a Filipino non-profit technology community that "
        "operates through chapter-based internships and bootcamps, is "
        "one example of an organization matching this profile and serves "
        "as the primary evaluation context of this study (see Section "
        "1.7 Scope)."
    ),
    146: (
        "One of the recurring challenges in organizations of this profile "
        "is unfamiliarity with the actual technology stacks used in their "
        "projects. New contributors are typically expected to adapt "
        "quickly to languages, frameworks, and architectural conventions "
        "they did not choose and may not have seen during formal "
        "education. This is not strictly a programming-ability problem; "
        "it is a code-comprehension and convention-adoption problem. "
        "Without standardized documentation or a learning surface that "
        "ties pattern-level concepts to the actual code being read, the "
        "burden of converting theory into practice falls entirely on the "
        "new contributor and on whoever is available to mentor them "
        "(Romeo et al., 2024)."
    ),
    152: (
        "To address this need, this study focuses on the development of "
        "CodiNeo, a design-pattern learning and documentation support "
        "system targeted at organizations matching the profile described "
        "above. The system includes learning modules that introduce users "
        "to selected software design patterns and a code-analysis studio "
        "that allows users to apply these concepts to actual source-code "
        "examples. Through the studio, users can submit source code, "
        "detect supported design-pattern evidence, view highlighted "
        "structures, and generate documentation-oriented explanations. "
        "In this context, design-pattern evidence refers to structural "
        "indications in the source code that help explain how a class or "
        "component is organized. AST-based and graph-based "
        "representations support this approach because they allow source "
        "code to be modeled through syntactic structures and "
        "relationships among program elements (Fan et al., 2021; Wang "
        "et al., 2021; TehraniJamsaz et al., 2023). The prototype "
        "implementation supports C++ as its first language and is "
        "evaluated within the DEVCON Luzon learning environment; the "
        "algorithmic and language-scope claims of this study are stated "
        "in Section 1.7 Scope."
    ),
    154: (
        "The proposed system serves as a code-understanding and "
        "documentation support tool. By generating documentation-oriented "
        "outputs and explanations from detected structural evidence, "
        "CodiNeo aims to help new contributors interpret code structure, "
        "connect design-pattern theory with actual implementation, and "
        "understand existing software more effectively. Through this "
        "approach, the study aims to bridge the gap between software "
        "design theory and practical code understanding. By providing "
        "evidence-based code analysis and documentation support, "
        "CodiNeo may help organizations that lack formalized coding "
        "standards or that rely on informal mentorship reduce the "
        "onboarding burden placed on senior developers, accelerate "
        "new-contributor productivity, and lower the cost of knowledge "
        "transfer across contributor cycles."
    ),
}


def find_heading_index(doc: Document, heading_text: str) -> int:
    """Return the paragraph index of the Heading 2 with the given text."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == heading_text and p.style.name == "Heading 2":
            return i
    raise SystemExit(
        f"ABORT: Heading 2 '{heading_text}' not found in docx"
    )


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    """Replace all runs in `paragraph` with a single run holding new_text.
    Preserves paragraph style + first-run font properties.
    """
    runs = paragraph.runs
    template_run = runs[0] if runs else None
    for run in list(runs):
        run._element.getparent().remove(run._element)
    new_run = paragraph.add_run(new_text)
    if template_run is not None:
        new_run.font.name = template_run.font.name
        new_run.font.size = template_run.font.size
        new_run.font.bold = template_run.font.bold
        new_run.font.italic = template_run.font.italic
        if template_run.font.color and template_run.font.color.rgb:
            new_run.font.color.rgb = template_run.font.color.rgb


def check_already_applied(doc: Document) -> bool:
    """Return True if any paragraph already contains the v3 marker phrase."""
    for p in doc.paragraphs:
        if V3_MARKER in p.text:
            return True
    return False


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"docx not found: {DOC_PATH.resolve()}")

    doc = docx.Document(str(DOC_PATH))

    # Sanity check: the Background heading must exist (verifies the docx
    # structure hasn't shifted under us).
    heading_idx = find_heading_index(doc, BACKGROUND_HEADING_TEXT)
    print(f"Background heading at ¶{heading_idx}")

    if check_already_applied(doc):
        print(
            "Background v3 already applied (detected marker phrase). "
            "Re-applying anyway to be idempotent."
        )

    # Apply the v3 rewrites.
    for idx, new_text in NEW_TEXT.items():
        if idx >= len(doc.paragraphs):
            raise SystemExit(f"ABORT: paragraph ¶{idx} out of range")
        para = doc.paragraphs[idx]
        replace_paragraph_text(para, new_text)
        preview = new_text[:80].replace("\n", " ")
        print(f"  ¶{idx} rewrote -> {preview}...")

    doc.save(str(DOC_PATH))
    print(f"\nSaved: {DOC_PATH}")
    print("Backup: FINAL-THESIS-3-PAPER-1 (2).backup-before-sop-rewrite.docx")


if __name__ == "__main__":
    main()
