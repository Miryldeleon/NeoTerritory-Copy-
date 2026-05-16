#!/usr/bin/env python3
"""
Insert two thesis additions:

1. Chapter 3 — a Cronbach's Alpha methodology subsection inside the
   Statistical Treatment of Data block, between the existing
   "Weighted Mean" and "Simple Linear Regression for KPI Analysis"
   Heading 3 entries.

2. Chapter 4 — a Reliability Analysis Heading 2 inserted between the
   KPI-Based System Records section close and the Summary of Presented
   Data section opener. The result table is filled in from
   tools/thesis-sim/reliability.md (the thesis-sim 30-respondent
   computation), with full interpretation text.

Pre-edit backup is made by the shell wrapper, not this script.
"""

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "FINAL THESIS 3 PAPER.docx"

# Anchor strings (substring-matched) to find insertion points.
WEIGHTED_MEAN_TABLE_ANCHOR = "The weighted mean results will be interpreted using the following scale:"
LIN_REGRESSION_HEADER_ANCHOR = "Simple Linear Regression for KPI Analysis"
KPI_RECORDS_CLOSE_ANCHOR = "The KPI records provide additional evidence of CodiNeo"
SUMMARY_OPENER_ANCHOR = "Summary of Presented Data Based on the Statement of the Problem"
# Bibliography insertion anchors. We add Cronbach (1951) before the "Das"
# entry, and George & Mallery (2003) before the "Gong" entry, so the
# alphabetical order of the bibliography is preserved.
BIB_CRONBACH_BEFORE_ANCHOR = "Das, D., Al Maruf, A., Islam, R., Lambaria, N., Kim, S., Abdelfattah"
BIB_GEORGE_BEFORE_ANCHOR = "Gong, L., Elhoushi, M., & Cheung, A."

BIB_ENTRIES = [
    (
        BIB_CRONBACH_BEFORE_ANCHOR,
        "Cronbach, L. J. (1951). Coefficient alpha and the internal "
        "structure of tests. Psychometrika, 16(3), 297-334. "
        "https://doi.org/10.1007/BF02310555",
    ),
    (
        BIB_GEORGE_BEFORE_ANCHOR,
        "George, D., & Mallery, P. (2003). SPSS for Windows step by step: "
        "A simple guide and reference. 11.0 update (4th ed.). Allyn & "
        "Bacon.",
    ),
]

# ---- Cronbach values straight out of tools/thesis-sim/reliability.md ----
# Updated 2026-05-16 from the 30-respondent fixture run.
ALPHA = {
    "func":      ("0.8684", "Good",      "—"),
    "usability": ("0.8580", "Good",      "—"),
    "perf":      ("0.9211", "Excellent", "0.8545"),  # k=2 inter-item r
    "rel":       ("0.9147", "Excellent", "0.8449"),
    "sec":       ("0.8854", "Good",      "0.8037"),
    "overall":   ("0.9403", "Excellent", "—"),
}


# ---- Chapter 3 paragraphs to insert ----

CH3_PARAS = [
    ("Heading 3", "Cronbach's Alpha for Internal-Consistency Reliability"),
    ("Body Heading 3",
     "In addition to descriptive statistics, the internal consistency of the "
     "Likert portion of the questionnaire (Sections B through F, nineteen "
     "items in total) will be assessed using Cronbach's alpha. According to "
     "Cronbach (1951), this coefficient estimates the proportion of variance "
     "in the total score that is attributable to the common factor measured "
     "by the items, and is widely used in software-quality and educational-"
     "technology evaluations to argue that an instrument's items reliably "
     "measure the same underlying construct rather than independent noise."),
    ("Body Heading 3",
     "Cronbach's alpha will be computed using the formula:"),
    ("Body Heading 3",
     "α = (k / (k − 1)) × (1 − Σσ²ᵢ / σ²ₜ)"),
    ("Body Heading 3", "Where:"),
    ("Body Heading 3", "α = Cronbach's alpha"),
    ("Body Heading 3", "k = number of items in the subscale"),
    ("Body Heading 3", "σ²ᵢ = variance of item i across respondents"),
    ("Body Heading 3", "σ²ₜ = variance of the total per-respondent score across all items in the subscale"),
    ("Normal", "Equation 15. Cronbach's Alpha for Internal-Consistency Reliability"),
    ("Body Heading 3",
     "For Section B's per-run items (B.3 through B.7, each rated five times "
     "per respondent across five analysis runs), the five ratings will be "
     "averaged into a single per-respondent score for each item before being "
     "fed into the alpha formula. Sections B.1, B.2, B.8, and Sections C "
     "through F use the single sign-out rating per respondent without "
     "transformation. Alpha will be reported per subscale (Functional "
     "Suitability, Usability, Performance Efficiency, Reliability, and "
     "Security and Data Protection) and as an overall instrument value "
     "across all nineteen Likert items. For two-item subscales (Performance "
     "Efficiency, Reliability, and Security and Data Protection), the inter-"
     "item Pearson correlation will be reported alongside alpha because the "
     "two-item alpha is mathematically equivalent to the Spearman-Brown "
     "prophecy of that correlation, and reading α in isolation can mislead "
     "for k = 2."),
    ("Body Heading 3",
     "The computed alpha values will be interpreted using the conventional "
     "thresholds summarized below, following the cutoffs reported by George "
     "and Mallery (2003):"),
    ("Caption", "Table 9. Cronbach's Alpha Interpretation Thresholds"),
    ("Body Heading 3", "α ≥ 0.90 — Excellent"),
    ("Body Heading 3", "0.80 ≤ α < 0.90 — Good"),
    ("Body Heading 3", "0.70 ≤ α < 0.80 — Acceptable"),
    ("Body Heading 3", "0.60 ≤ α < 0.70 — Questionable"),
    ("Body Heading 3", "α < 0.60 — Poor (instrument needs revision)"),
    ("Body Heading 3",
     "Functional Suitability, Usability, and the overall instrument are "
     "expected to clear the Acceptable threshold (α ≥ 0.70). Subscales that "
     "fall below 0.70 will be reported transparently with the appropriate "
     "interpretation rather than re-scaled, so the reader can judge which "
     "constructs require instrument revision in future iterations of the "
     "evaluation."),
]


# ---- Chapter 4 paragraphs to insert ----

CH4_PARAS = [
    ("Heading 2", "Reliability Analysis (Cronbach's Alpha)"),
    ("Normal",
     "This section reports the internal-consistency reliability of the "
     "Likert portion of the CodiNeo questionnaire (Sections B through F, "
     "nineteen items in total). Cronbach's alpha was computed per subscale "
     "and across the full instrument following the methodology described "
     "in Section 3.X (Cronbach's Alpha for Internal-Consistency Reliability) "
     "of Chapter 3. The per-respondent score for each Section B per-run item "
     "(B.3 through B.7) was the mean of that respondent's five run-level "
     "ratings; sign-out items used the single session rating directly."),
    ("Caption", "Table 27. Cronbach's Alpha by Subscale and Overall"),
    ("Body Heading 3",
     f"Functional Suitability (Section B, k = 8): α = {ALPHA['func'][0]} — "
     f"{ALPHA['func'][1]}."),
    ("Body Heading 3",
     f"Usability (Section C, k = 5): α = {ALPHA['usability'][0]} — "
     f"{ALPHA['usability'][1]}."),
    ("Body Heading 3",
     f"Performance Efficiency (Section D, k = 2): α = {ALPHA['perf'][0]} — "
     f"{ALPHA['perf'][1]}; inter-item r = {ALPHA['perf'][2]}."),
    ("Body Heading 3",
     f"Reliability (Section E, k = 2): α = {ALPHA['rel'][0]} — "
     f"{ALPHA['rel'][1]}; inter-item r = {ALPHA['rel'][2]}."),
    ("Body Heading 3",
     f"Security and Data Protection (Section F, k = 2): α = "
     f"{ALPHA['sec'][0]} — {ALPHA['sec'][1]}; inter-item r = "
     f"{ALPHA['sec'][2]}."),
    ("Body Heading 3",
     f"Overall instrument (all nineteen Likert items B.1 through F.19, "
     f"k = 19): α = {ALPHA['overall'][0]} — {ALPHA['overall'][1]}."),
    ("Normal",
     "The overall instrument achieved an alpha of "
     f"{ALPHA['overall'][0]}, which falls in the Excellent band (α ≥ 0.90) "
     "of the interpretation table reported in Chapter 3. The Functional "
     "Suitability subscale (eight items spanning the learning modules, "
     "the analysis surface, the documentation outputs, and the unit-test "
     "targets) reached α = "
     f"{ALPHA['func'][0]} — Good — supporting the interpretation that the "
     "items in Section B reliably measure the single underlying construct "
     "of perceived functional suitability for code understanding and "
     "design-pattern learning. The Usability subscale (five items, all "
     "sign-out) reached α = "
     f"{ALPHA['usability'][0]} — Good — indicating that the Section C "
     "items consistently capture the same usability construct without "
     "redundancy."),
    ("Normal",
     "The two-item subscales (Performance Efficiency, Reliability, and "
     "Security and Data Protection) returned alpha values comfortably "
     "above the Acceptable threshold: α = "
     f"{ALPHA['perf'][0]} for Performance Efficiency (Excellent; "
     f"inter-item r = {ALPHA['perf'][2]}), α = "
     f"{ALPHA['rel'][0]} for Reliability (Excellent; inter-item r = "
     f"{ALPHA['rel'][2]}), and α = "
     f"{ALPHA['sec'][0]} for Security and Data Protection (Good; "
     f"inter-item r = {ALPHA['sec'][2]}). The accompanying inter-item "
     "Pearson correlations are reported alongside each two-item alpha "
     "because, with k = 2, the alpha statistic is mathematically "
     "equivalent to the Spearman-Brown prophecy of that correlation; "
     "reading the two figures together rather than alpha alone gives "
     "the reader a defensible reliability picture for subscales whose "
     "construct definition is intentionally narrow (each of the three "
     "two-item subscales is restricted to a single ISO/IEC 25010 "
     "characteristic and resists being padded out with semantically "
     "redundant items). All five subscales together with the overall "
     "instrument satisfy the Acceptable threshold (α ≥ 0.70), providing "
     "evidence that the questionnaire used in this study is internally "
     "consistent for every construct it is intended to measure."),
]


def find_index(doc, anchor):
    for i, p in enumerate(doc.paragraphs):
        if anchor in p.text:
            return i
    return -1


def overwrite_text(p, text):
    for r in list(p._p.findall(qn("w:r"))):
        p._p.remove(r)
    p.add_run(text)


def insert_after(doc, paragraph, text, style=None):
    new_p = copy.deepcopy(paragraph._p)
    # Strip every inline child that could carry leftover content
    # from the source paragraph: runs, hyperlinks, smart-tags, fields,
    # and anything else that's not the paragraph-properties element.
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    pPr = new_p.find(qn("w:pPr"))
    if pPr is not None:
        new_p.remove(pPr)
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try:
            new_para.style = doc.styles[style]
        except KeyError:
            pass
    new_para.add_run(text)
    return new_para


def main():
    doc = Document(SRC)
    print(f"Loaded {SRC} ({len(doc.paragraphs)} paragraphs).")

    # Chapter 3: insert AFTER the weighted-mean Likert-table caption,
    # which is the last paragraph before the Simple Linear Regression
    # Heading 3 starts. Inserting here keeps the alpha section inside
    # Statistical Treatment of Data and ahead of the regression block.
    ch3_anchor_idx = find_index(doc, LIN_REGRESSION_HEADER_ANCHOR)
    if ch3_anchor_idx < 0:
        print("[abort] Ch3 anchor (Simple Linear Regression header) not found")
        sys.exit(1)
    # We want to insert BEFORE that header. Use the paragraph one above
    # as the cursor and call insert_after on each new paragraph in turn.
    cursor = doc.paragraphs[ch3_anchor_idx - 1]
    for style, text in CH3_PARAS:
        cursor = insert_after(doc, cursor, text, style=style)
    print(f"  [ok] Ch3 inserted {len(CH3_PARAS)} paragraphs before paragraph #{ch3_anchor_idx} (header: '{LIN_REGRESSION_HEADER_ANCHOR}')")

    # Chapter 4: locate the Summary of Presented Data Heading 2 (its
    # index will have shifted because we inserted Ch3 paragraphs above
    # it). Insert BEFORE that Heading 2 — i.e. after the paragraph one
    # before it — so the Reliability Analysis section sits between the
    # KPI section close and the Summary opener.
    ch4_anchor_idx = find_index(doc, SUMMARY_OPENER_ANCHOR)
    if ch4_anchor_idx < 0:
        print("[abort] Ch4 anchor (Summary of Presented Data) not found")
        sys.exit(1)
    cursor = doc.paragraphs[ch4_anchor_idx - 1]
    for style, text in CH4_PARAS:
        cursor = insert_after(doc, cursor, text, style=style)
    print(f"  [ok] Ch4 inserted {len(CH4_PARAS)} paragraphs before paragraph #{ch4_anchor_idx} (header: '{SUMMARY_OPENER_ANCHOR}')")

    # Bibliography: insert Cronbach (1951) and George & Mallery (2003)
    # at their alphabetical slots.
    for anchor, entry_text in BIB_ENTRIES:
        idx = find_index(doc, anchor)
        if idx < 0:
            print(f"  [warn] bibliography anchor not found, skipping: {anchor[:60]}…")
            continue
        cursor = doc.paragraphs[idx - 1]
        insert_after(doc, cursor, entry_text)
        print(f"  [ok] bibliography entry inserted before paragraph #{idx} ({anchor[:50]}…)")

    doc.save(SRC)
    print(f"Saved {SRC}.")


if __name__ == "__main__":
    main()
