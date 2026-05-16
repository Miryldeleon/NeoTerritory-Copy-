#!/usr/bin/env python3
"""
Insert best-case / average-case / worst-case paragraphs into Sections
3.4.1 (Time Complexity Analysis) and 3.4.2 (Space Complexity) of the
FINAL THESIS 3 PAPER.docx.

We anchor the insertion to specific concluding sentences in each section
so the new paragraphs land at the end of the section, before the next
Heading3 starts.
"""

import copy
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "FINAL THESIS 3 PAPER.docx"
BACKUP_GLOB = "FINAL THESIS 3 PAPER.backup-pre-bawc-*.docx"

# Anchor text — last sentence of Section 3.4.1 (time complexity).
TIME_ANCHOR = (
    "The AI-assisted explanation layer and testing/GDB runner layer are not "
    "included in the core time complexity"
)
# Anchor text — last sentence of Section 3.4.2 (space complexity).
SPACE_ANCHOR = (
    "the system’s space usage remains mainly proportional to the number of "
    "relevant structural elements"
)

# What we want to insert. Each list element is (style_name, text).
# style_name == "BodyHeading3" so it matches the thesis's own subhead.
TIME_BAWC_PARAS = [
    ("BodyHeading3", "Best-Case, Average-Case, and Worst-Case Time Behavior"),
    (
        None,
        "The asymptotic bound derived above describes the dominant cost as "
        "the submitted input grows. To make the analysis more useful for "
        "implementers and reviewers, the behavior of the algorithm is "
        "further characterized under three input regimes that the prototype "
        "is expected to encounter in practice: a best-case input, an "
        "average-case input representative of a typical DEVCON Luzon intern "
        "submission, and a worst-case input that approaches the upper "
        "bound of what the supported prototype is designed to accept.",
    ),
    (
        None,
        "Best case. The best-case input is a single short C++ translation "
        "unit containing one fully defined class or struct, no inheritance "
        "relationships, and a small number of methods or attributes whose "
        "token signatures do not match any pattern in the catalog beyond "
        "the cheapest first check. In this regime every stage of the "
        "pipeline still executes, but the per-stage work collapses to a "
        "near-constant cost: lexical scanning produces a small token "
        "stream, the class-level structural representation contains a "
        "single entry, the pattern dispatch stage prunes each pattern "
        "after its first ordered check fails, hashing performs a constant "
        "number of average-case O(1) inserts, and the report writer emits "
        "a short JSON payload. The total work is bounded by O(L) where L "
        "is small and dominated by the fixed startup cost of loading the "
        "pattern catalog from disk; observed running time approaches a "
        "constant floor as L shrinks.",
    ),
    (
        None,
        "Average case. The average-case input is the kind of submission an "
        "intern is expected to produce during the evaluation cycle: a "
        "small set of files containing a handful of related classes that "
        "exhibit one or two of the supported design-pattern shapes "
        "(Singleton, Builder, Factory, Strategy, Observer, Composite, "
        "Proxy, or Method Chaining), with method bodies of moderate size. "
        "Under this regime each stage performs work proportional to the "
        "total number of relevant tokens and structural elements, the "
        "pattern dispatch stage rejects most candidate patterns early "
        "while accepting a small number per class, hashing remains O(1) "
        "on average, and the report writer emits a payload whose size is "
        "proportional to the number of detected patterns. The empirical "
        "regression conducted in Section 3.4.4 against synthetic inputs "
        "ranging from approximately 100 to 20,000 lines of C++ confirmed "
        "that the wall-clock running time fits a linear model with a "
        "coefficient of determination of approximately 0.96, which is "
        "consistent with the linear average-case prediction.",
    ),
    (
        None,
        "Worst case. The worst-case input is the upper end of what the "
        "supported prototype is designed to accept: a maximum-size "
        "submission within the configured per-file token cap, a class "
        "count near the structural-element ceiling enforced by the "
        "analyzer, and a token distribution in which every class triggers "
        "a candidate match for every pattern in the catalog (so the early "
        "rejection optimization of the pattern dispatch stage is "
        "minimized). Under this regime every supported pattern check is "
        "evaluated against every class, and the tag-construction step "
        "produces an output record for each accepted match. With the "
        "number of patterns P treated as a fixed constant under the "
        "prototype’s catalog, the work is still bounded by the same "
        "asymptotic class O(L); the practical effect of the worst case "
        "is a larger constant factor and a higher per-class evaluation "
        "cost, not a change in the order of growth. This is the regime in "
        "which empirical measurement is most useful, because it exposes "
        "whether the constant factor remains acceptable for the "
        "submission sizes the system will see in real use.",
    ),
    (
        None,
        "Across all three regimes the asymptotic bound stated earlier "
        "continues to hold under the prototype’s assumptions that the "
        "pattern catalog size P is fixed and that the count of structural "
        "elements E is bounded by the size of the submitted source. The "
        "three regimes differ in their constant factor, not in their "
        "order of growth, and the empirical regression supports this "
        "conclusion within the size range the system is expected to be "
        "used in.",
    ),
]

SPACE_BAWC_PARAS = [
    ("BodyHeading3", "Best-Case, Average-Case, and Worst-Case Space Behavior"),
    (
        None,
        "As with time complexity, the space behavior of the analysis "
        "pipeline is characterized under the same three input regimes: a "
        "best-case input, an average-case input representative of a "
        "typical intern submission, and a worst-case input near the "
        "supported upper bound. The three regimes share the same "
        "asymptotic class established above; they differ in the size of "
        "the constant factor and in which retained data structure "
        "dominates the residency footprint at run time.",
    ),
    (
        None,
        "Best case. For a short single-class input the retained data "
        "structures are dominated by the loaded pattern catalog, which is "
        "fixed in size and treated as a constant. The lexical token "
        "buffer, the per-class structural representation, the hash-link "
        "table, and the output buffer each hold only a small number of "
        "entries. Total peak memory approaches a constant floor "
        "determined by the catalog and the static segment of the binary; "
        "observed peak working-set grows only marginally with input size "
        "in this regime.",
    ),
    (
        None,
        "Average case. For a typical intern submission the retained data "
        "structures scale linearly with the number of relevant tokens, "
        "structural elements, and detected matches. The lexical token "
        "buffer, the per-class structural representation, and the hash "
        "index together account for most of the peak working set, while "
        "the output buffer adds a contribution proportional to the number "
        "of detected patterns. The empirical regression in Section 3.4.4 "
        "fits peak working-set memory to a linear model in the synthetic "
        "input size with a coefficient of determination of approximately "
        "0.99, which is consistent with the linear average-case "
        "prediction.",
    ),
    (
        None,
        "Worst case. For a maximum-size submission in which every class "
        "triggers a match for every supported pattern, the dominant "
        "retained structure is the set of accepted pattern matches and "
        "the per-class token streams those matches reference. Because the "
        "number of patterns is fixed, the worst-case retained footprint "
        "is bounded by the same linear class as the average case, but "
        "with a noticeably larger constant factor. Transient intermediate "
        "structures during the pattern dispatch stage (the candidate-"
        "match list before filtering and ranking) reach their largest "
        "size in this regime; they are released before the output stage "
        "executes, so they affect peak working set but not the long-lived "
        "residency.",
    ),
    (
        None,
        "Across all three regimes the space behavior remains within the "
        "linear bound stated above under the prototype’s assumptions on "
        "P and E. The empirical fit on synthetic inputs from "
        "approximately 100 to 20,000 lines confirms that peak memory "
        "grows linearly within the supported input range and that the "
        "constant factor remains within the host’s capacity for the "
        "submission sizes targeted by the system.",
    ),
]


def find_paragraph_index(doc, anchor_substring):
    for idx, p in enumerate(doc.paragraphs):
        if anchor_substring in p.text:
            return idx
    return -1


def insert_paragraph_after(doc, paragraph, text, style=None):
    new_p = copy.deepcopy(paragraph._p)
    for r in list(new_p.findall(qn("w:r"))):
        new_p.remove(r)
    pPr = new_p.find(qn("w:pPr"))
    if pPr is not None:
        new_p.remove(pPr)
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try:
            new_para.style = doc.styles[style]
        except KeyError:
            pass
    new_para.add_run(text)
    return new_para


def inject(doc, anchor, paras, label):
    idx = find_paragraph_index(doc, anchor)
    if idx < 0:
        raise RuntimeError(f"anchor for {label} not found: {anchor[:60]}…")
    anchor_para = doc.paragraphs[idx]
    cursor = anchor_para
    inserted = 0
    for style, text in paras:
        cursor = insert_paragraph_after(doc, cursor, text, style=style)
        inserted += 1
    print(f"  [{label}] anchor index={idx}, inserted {inserted} paragraphs after it.")


def main():
    doc = Document(SRC)
    print(f"Loaded {SRC} ({len(doc.paragraphs)} paragraphs).")
    inject(doc, TIME_ANCHOR, TIME_BAWC_PARAS, "time")
    inject(doc, SPACE_ANCHOR, SPACE_BAWC_PARAS, "space")
    doc.save(SRC)
    print(f"Saved updated {SRC}.")


if __name__ == "__main__":
    main()
